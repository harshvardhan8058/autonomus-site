"""Tests for generated-document quality (clean headings + meaningful bullets).

These example-based tests lock in two post-completion quality fixes:

- **Clean section headings** — when a plan's steps carry explicit
  ``section_title`` values, the produced ``.docx`` renders those exact titles as
  its section headings and lists them in the table of contents (rather than
  truncated, awkward phrases derived from the verbose task/description).
- **Meaningful bullet list** — when the Run recorded planner assumptions, the
  deliverable contains a dedicated "Key Assumptions" bullet section, and the
  synthesized fallback bullet list is the generic "Key Considerations" set
  rather than a duplicate of the section headings / table of contents.

The heading-derivation helpers (:func:`humanize_section_title`,
:func:`section_heading_for_step`) are also unit-tested directly to guarantee no
mid-phrase truncation.
"""

from __future__ import annotations

import asyncio
import json
from pathlib import Path

from docx import Document

from app.agent.executor import (
    Executor,
    humanize_section_title,
    section_heading_for_step,
)
from app.agent.orchestrator import Orchestrator
from app.agent.planner import Planner
from app.agent.reflector import Reflector
from app.agent.tools import build_default_registry
from app.core.config import Settings
from app.core.event_bus import EventBus
from app.core.run_store import RunStore
from app.models.schemas import Plan, PlanStep
from app.services.docx_builder import DocumentBuilder
from app.services.llm import LLMService
from tests.conftest import FakeLLMBackend


async def _no_sleep(_delay: float) -> None:
    """A no-op async sleep so retry backoff introduces no real delay."""

    return None


def _build_llm(plan_json: str) -> LLMService:
    """Build an :class:`LLMService` wired to fake backends returning ``plan_json``."""

    return LLMService(
        Settings(GROQ_API_KEY="test-key"),
        groq_backend=FakeLLMBackend("groq", response=plan_json),
        ollama_backend=FakeLLMBackend("ollama", response=plan_json),
        sleep=_no_sleep,
        resolve=True,
    )


def _plan_json() -> str:
    """A canned plan whose content steps carry explicit ``section_title`` values."""

    return json.dumps(
        {
            "steps": [
                {
                    "step": 1,
                    "task": "research",
                    "section_title": "Cloud CRM Overview",
                    "description": (
                        "Research the introduction and executive summary "
                        "sections of the CRM landscape."
                    ),
                    "expected_output": "Researched facts.",
                },
                {
                    "step": 2,
                    "task": "draft_section",
                    "section_title": "On-Premise vs Cloud Comparison",
                    "description": "Draft the comparison of the two approaches.",
                    "expected_output": "A drafted section.",
                },
                {
                    "step": 3,
                    "task": "build_docx",
                    "section_title": "",
                    "description": "Assemble the final Word document.",
                    "expected_output": "The .docx deliverable.",
                },
            ],
            "assumptions": [
                "The target cloud is a major public cloud provider.",
                "The audience is executive leadership.",
            ],
        }
    )


def test_docx_headings_equal_section_titles_and_toc_lists_them(
    tmp_path: Path, monkeypatch,
) -> None:
    """The produced .docx headings equal the planner-provided section titles.

    Given a plan whose steps carry explicit ``section_title`` values, the built
    document renders those exact titles as headings and lists them in the table
    of contents; and because the Run recorded assumptions, a dedicated
    "Key Assumptions" bullet section is present (not a duplicate of the TOC).
    """

    # Build documents under a temporary working directory so the executor's
    # relative ``generated/`` output path does not pollute the repo.
    monkeypatch.chdir(tmp_path)

    async def _run() -> Path:
        llm = _build_llm(_plan_json())
        doc_builder = DocumentBuilder("1F4E79")
        bus = EventBus()
        store = RunStore()
        registry = build_default_registry(llm, doc_builder)
        planner = Planner(llm)
        executor = Executor(registry, bus, max_retries=0)
        reflector = Reflector(llm, bus)
        orchestrator = Orchestrator(
            validator=None,  # type: ignore[arg-type] - unused in execute_run
            planner=planner,
            executor=executor,
            reflector=reflector,
            doc_builder=doc_builder,
            store=store,
            events=bus,
        )
        run_state = store.create(
            "quality-run",
            request="Create a CRM cloud-migration proposal.",
            client_ip="127.0.0.1",
        )
        await orchestrator.execute_run(run_state)
        assert run_state.document_path is not None
        return run_state.document_path

    document_path = asyncio.run(_run())
    assert document_path.exists()

    document = Document(str(document_path))
    heading_texts = [
        p.text
        for p in document.paragraphs
        if p.style is not None and p.style.name.startswith("Heading")
    ]

    # The clean, planner-provided titles appear verbatim as document headings.
    assert "Cloud CRM Overview" in heading_texts
    assert "On-Premise vs Cloud Comparison" in heading_texts

    # A synthesized "Executive Summary" is now the FIRST content heading (the
    # "Table of Contents" heading precedes all content headings but is not one).
    assert "Executive Summary" in heading_texts
    content_headings = [h for h in heading_texts if h != "Table of Contents"]
    assert content_headings[0] == "Executive Summary"

    # No awkward, truncated verb-led phrase leaks into the headings.
    assert not any(h.startswith("Research The") for h in heading_texts)
    assert not any(h.startswith("Draft The") for h in heading_texts)

    # A dedicated "Key Assumptions" bullet section is present with real bullets.
    assert "Key Assumptions" in heading_texts
    bullet_texts = [
        p.text
        for p in document.paragraphs
        if p.style is not None and p.style.name == "List Bullet"
    ]
    assert "The target cloud is a major public cloud provider." in bullet_texts
    assert "The audience is executive leadership." in bullet_texts

    # The table of contents lists the real section headings (each heading appears
    # both as its styled heading and as a visible TOC entry). TOC entry text now
    # also carries a tab and a page number, so use substring (not exact-count)
    # matching.
    paragraph_texts = [p.text for p in document.paragraphs]
    assert sum(1 for t in paragraph_texts if "Cloud CRM Overview" in t) >= 2
    assert sum(1 for t in paragraph_texts if "On-Premise vs Cloud Comparison" in t) >= 2


class _CapturingBuilder:
    """A minimal document builder that records the sections it was asked to build.

    It writes a tiny placeholder file so the Executor can record a document path,
    and captures the exact ordered ``sections`` list passed to :meth:`build` so a
    test can assert the assembly order (Executive Summary first, Key Assumptions
    last) without parsing a real ``.docx``.
    """

    def __init__(self) -> None:
        self.sections: list[dict] | None = None

    def build(self, *, title, prepared_by, sections, output_path) -> Path:
        self.sections = [dict(section) for section in sections]
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(b"PK\x03\x04 fake-docx")
        return output_path


def test_executor_synthesizes_executive_summary_first_and_assumptions_last(
    tmp_path: Path, monkeypatch,
) -> None:
    """The Executor inserts a synthesized Executive Summary first (Fix 2).

    Driving the real Executor (with a fake LLM) over a plan with two content
    steps and a ``build_docx`` step, the sections handed to ``build_docx`` must
    begin with a synthesized "Executive Summary" and end with "Key Assumptions"
    (present because the Run recorded assumptions).
    """

    monkeypatch.chdir(tmp_path)

    async def _run() -> list[dict]:
        llm = _build_llm(_plan_json())
        builder = _CapturingBuilder()
        bus = EventBus()
        store = RunStore()
        registry = build_default_registry(llm, builder)
        executor = Executor(registry, bus, max_retries=0)
        run_state = store.create(
            "exec-summary-run",
            request="Create a CRM cloud-migration proposal.",
            client_ip="127.0.0.1",
        )
        run_state.plan = Plan(
            steps=[
                PlanStep(
                    step=1,
                    task="research",
                    section_title="Cloud CRM Overview",
                    description="Research the CRM landscape.",
                    expected_output="Researched facts.",
                ),
                PlanStep(
                    step=2,
                    task="draft_section",
                    section_title="On-Premise vs Cloud Comparison",
                    description="Draft the comparison of the two approaches.",
                    expected_output="A drafted section.",
                ),
                PlanStep(
                    step=3,
                    task="build_docx",
                    section_title="",
                    description="Assemble the final Word document.",
                    expected_output="The .docx deliverable.",
                ),
            ],
            assumptions=[],
        )
        run_state.assumptions = [
            "The target cloud is a major public cloud provider.",
            "The audience is executive leadership.",
        ]
        await executor.run(run_state)
        assert builder.sections is not None
        return builder.sections

    sections = asyncio.run(_run())

    # The synthesized executive summary is first and has a non-empty body.
    assert sections[0]["heading"] == "Executive Summary"
    assert sections[0].get("body", "").strip()

    # The two content sections are present between the summary and assumptions.
    headings = [section["heading"] for section in sections]
    assert "Cloud CRM Overview" in headings
    assert "On-Premise vs Cloud Comparison" in headings

    # The Key Assumptions bullet section remains last.
    assert sections[-1]["heading"] == "Key Assumptions"
    assert headings.index("Executive Summary") == 0
    assert headings.index("Key Assumptions") == len(headings) - 1


# ---------------------------------------------------------------------------
# Unit tests for heading derivation (no mid-phrase truncation)
# ---------------------------------------------------------------------------


def test_section_heading_prefers_explicit_section_title() -> None:
    """An explicit ``section_title`` is used verbatim over any derivation."""

    step = PlanStep(
        step=1,
        task="draft_section",
        section_title="Executive Summary",
        description="Write the introduction and executive summary sections of.",
        expected_output="x",
    )
    assert section_heading_for_step(step, 1) == "Executive Summary"


def test_humanize_strips_leading_verb_and_trailing_filler() -> None:
    """Deriving from a description yields a clean noun phrase, not a truncation."""

    heading = humanize_section_title(
        "draft_section",
        "Write the introduction and executive summary sections of the document",
        1,
    )
    assert heading == "Introduction And Executive Summary"


def test_humanize_does_not_cut_words_in_half() -> None:
    """A long description is capped at a whole-word boundary, never mid-word."""

    description = (
        "Create a detailed comparison of on-premise versus cloud-based customer "
        "relationship management platforms for enterprise organizations worldwide"
    )
    heading = humanize_section_title("generate_table_data", description, 2)
    assert len(heading) <= 60
    # Every rendered token is a complete word from the source (no partial words).
    source_words = {w.lower() for w in description.replace("-", " ").split()}
    for token in heading.replace("-", " ").split():
        assert token.lower() in source_words


def test_humanize_falls_back_to_positional_section() -> None:
    """With no usable task or description, a positional fallback is returned."""

    assert humanize_section_title("build_docx", "", 4) == "Section 4"
