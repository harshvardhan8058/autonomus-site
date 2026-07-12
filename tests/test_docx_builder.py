"""Property-based tests for :class:`app.services.docx_builder.DocumentBuilder`.

These tests validate two design correctness properties for the document builder:

- **Property 13** (Req 10.1): every generated ``.docx`` re-opens successfully
  with ``python-docx`` and contains the document title, a table-of-contents
  field element, at least one table, at least one bullet-list paragraph, styled
  headings, and a footer page-number (``PAGE``) field.
- **Property 15** (Req 14.3, 14.4, 14.5): theme-color resolution is safe and
  total -- for any ``THEME_COLOR`` string the resolved heading color equals the
  value when it is a valid 6-digit hex, otherwise the documented default;
  building always completes without raising, and an invalid value produces a
  structured warning.

All property tests use Hypothesis with a minimum of 100 iterations. Document
writing can be slower than pure-logic checks, so a ``deadline=None`` is used to
avoid spurious per-example timeouts.
"""

from __future__ import annotations

import json

import pytest
from docx import Document
from docx.shared import RGBColor
from hypothesis import HealthCheck, given, settings
from hypothesis import strategies as st

from app.core.logging import StructuredLogger
from app.services.docx_builder import DEFAULT_THEME_COLOR, DocumentBuilder

# --- Strategies -------------------------------------------------------------

# Text safe for XML serialization: excludes control characters (Cc) and
# surrogates (Cs) that lxml rejects, while still exercising a wide input space.
_safe_text = st.text(
    alphabet=st.characters(blacklist_categories=("Cc", "Cs")),
    max_size=40,
)

# Non-empty safe text for the document title so it can be located after reopen.
_title_text = st.text(
    alphabet=st.characters(min_codepoint=0x20, max_codepoint=0x7E),
    min_size=1,
    max_size=40,
).filter(lambda s: s.strip() != "")


@st.composite
def _sections(draw: st.DrawFn) -> list[dict]:
    """Generate an arbitrary list of section payloads.

    Sections may omit any key and may or may not carry bullet lists and tables,
    exercising the builder's tolerance of missing keys and its synthesis of
    required structural elements when inputs do not provide them.
    """

    def _table() -> st.SearchStrategy:
        return st.builds(
            lambda headers, rows: {"headers": headers, "rows": rows},
            headers=st.lists(_safe_text, min_size=1, max_size=4),
            rows=st.lists(
                st.lists(_safe_text, min_size=1, max_size=4),
                max_size=4,
            ),
        )

    section_strategy = st.fixed_dictionaries(
        {},
        optional={
            "heading": _safe_text,
            "level": st.sampled_from([1, 2, 3, 0]),
            "body": _safe_text,
            "bullets": st.lists(_safe_text, max_size=4),
            "table": _table(),
        },
    )
    return draw(st.lists(section_strategy, max_size=4))


def _section_supplies_table(section: dict) -> bool:
    """Return whether ``section`` supplies real tabular data.

    Mirrors the builder's rule for rendering a table: a valid ``table`` payload
    (non-empty ``headers`` and ``rows``) or a markdown pipe-table in the body (a
    header row immediately followed by a dashes/colons separator row). No table
    is fabricated, so the built document contains a table only when this holds
    for at least one section.
    """

    table = section.get("table")
    if isinstance(table, dict):
        headers = table.get("headers")
        rows = table.get("rows")
        if (
            isinstance(headers, (list, tuple))
            and headers
            and isinstance(rows, (list, tuple))
            and rows
        ):
            return True

    body = section.get("body")
    if isinstance(body, str) and body:
        import re as _re

        separator = _re.compile(r"^:?-{1,}:?$")
        lines = [line.strip() for line in body.split("\n")]
        for index in range(len(lines) - 1):
            first, second = lines[index], lines[index + 1]
            if not (first.startswith("|") and first.endswith("|")):
                continue
            if not (second.startswith("|") and second.endswith("|")):
                continue
            header_cells = [c.strip() for c in first.split("|")[1:-1]]
            sep_cells = [c.strip() for c in second.split("|")[1:-1]]
            if (
                header_cells
                and sep_cells
                and all(separator.fullmatch(cell) is not None for cell in sep_cells)
            ):
                return True
    return False


def _valid_hex() -> st.SearchStrategy[str]:
    """Generate valid 6-digit hex colors, some with a leading ``#``."""

    base = st.text(alphabet="0123456789abcdefABCDEF", min_size=6, max_size=6)
    return st.one_of(base, base.map(lambda s: "#" + s))


def _expected_resolved(value: str) -> str:
    """Compute the expected resolved color independently of the builder.

    Mirrors the documented resolution rule: strip whitespace and one optional
    leading ``#``; if the remainder is a 6-digit hex string, its uppercase form
    is the resolved color; otherwise the documented default applies.
    """

    candidate = value.strip()
    if candidate.startswith("#"):
        candidate = candidate[1:]
    if len(candidate) == 6 and all(c in "0123456789abcdefABCDEF" for c in candidate):
        return candidate.upper()
    return DEFAULT_THEME_COLOR


# --- Property 13 ------------------------------------------------------------


# Feature: autonomous-agent-service, Property 13: Generated documents always parse and contain required structure  # noqa: E501
@pytest.mark.property
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture],
)
@given(title=_title_text, sections=_sections())
def test_generated_documents_parse_and_contain_required_structure(
    tmp_path_factory,
    title: str,
    sections: list[dict],
) -> None:
    """Property 13: generated documents parse and contain required structure.

    **Validates: Requirements 10.1**

    For arbitrary section inputs, the ``.docx`` produced by
    :class:`DocumentBuilder` re-opens successfully with ``python-docx`` and
    contains the document title, a TOC field element, at least one bullet-list
    paragraph, styled headings, and a footer ``PAGE`` field. A table is present
    only when at least one section supplies real tabular data; no table is
    fabricated, so a document may legitimately contain zero tables.
    """

    output_path = tmp_path_factory.mktemp("docx") / "deliverable.docx"

    builder = DocumentBuilder("1F4E79")
    written = builder.build(
        title=title,
        prepared_by="Autonomous Agent Service",
        sections=sections,
        output_path=output_path,
    )

    assert written == output_path
    assert written.exists()

    # Re-open the produced document; this must not raise.
    document = Document(str(written))

    # Title text is present among the body paragraphs.
    paragraph_texts = [p.text for p in document.paragraphs]
    assert title in paragraph_texts

    # A table-of-contents field element is present (the real TOC field
    # instruction still carries the ``TOC`` switch).
    document_xml = document.element.xml
    assert "TOC" in document_xml
    # The old, unhelpful placeholder is never emitted: the TOC now renders a
    # visible, readable list of headings instead (Req 10.1).
    assert "Right-click and choose" not in document_xml

    # A table is rendered only when a section supplies real tabular data; no
    # table is fabricated. When no section supplies one, zero tables is valid.
    expected_table = any(_section_supplies_table(section) for section in sections)
    if expected_table:
        assert len(document.tables) >= 1
    else:
        assert len(document.tables) == 0

    # At least one bullet-list paragraph exists.
    bullet_paragraphs = [
        p for p in document.paragraphs if p.style is not None and p.style.name == "List Bullet"
    ]
    assert len(bullet_paragraphs) >= 1

    # At least one styled heading exists.
    heading_paragraphs = [
        p
        for p in document.paragraphs
        if p.style is not None and p.style.name.startswith("Heading")
    ]
    assert len(heading_paragraphs) >= 1

    # The footer contains a page-number (PAGE) field.
    footer_xml = document.sections[0].footer._element.xml
    assert "PAGE" in footer_xml


# --- TOC rendering (unit) ---------------------------------------------------


def test_table_of_contents_lists_section_headings(tmp_path) -> None:
    """The TOC renders a visible, readable list of the section headings (Req 10.1).

    The table of contents must be visible in any viewer (not a placeholder that
    depends on the reader refreshing fields). The cached TOC result therefore
    contains one entry per document heading, so each heading appears both as its
    real styled heading AND as a TOC entry. The old "Right-click…" placeholder
    must be gone, and the field must remain a genuine Word ``TOC`` field.
    """

    builder = DocumentBuilder("1F4E79")
    sections = [
        {
            "heading": "Market Analysis Overview",
            "level": 1,
            "body": "Body text.",
            "bullets": ["A key point"],
            "table": {"headers": ["Metric"], "rows": [["Value"]]},
        },
        {"heading": "Detailed Findings", "level": 2, "body": "More detail."},
    ]

    output_path = tmp_path / "toc.docx"
    written = builder.build(
        title="Quarterly Report",
        prepared_by="Tester",
        sections=sections,
        output_path=output_path,
    )

    document = Document(str(written))
    document_xml = document.element.xml

    # The field is still a real Word TOC field, and the placeholder is gone.
    assert "TOC" in document_xml
    assert "Right-click and choose" not in document_xml

    # Each heading appears at least twice: once as the styled heading, once
    # inside a visible TOC entry. The TOC entry text now also carries a tab and a
    # page number, so use substring (not exact-equality) matching.
    paragraph_texts = [p.text for p in document.paragraphs]
    assert sum(1 for t in paragraph_texts if "Market Analysis Overview" in t) >= 2
    assert sum(1 for t in paragraph_texts if "Detailed Findings" in t) >= 2

    # The cached TOC entries are backed by PAGEREF fields (so Word fills in exact
    # pages on open) and show a visible page number with a dotted leader tab.
    assert "PAGEREF" in document_xml
    assert "w:tab" in document_xml
    # A visible cached page number appears in a TOC entry (content starts on
    # page 2, so the first entry shows "2").
    toc_entry_texts = [
        t
        for t in paragraph_texts
        if "Market Analysis Overview" in t or "Detailed Findings" in t
    ]
    assert any(any(ch.isdigit() for ch in t) for t in toc_entry_texts)


def test_no_table_is_fabricated_when_no_section_supplies_one(tmp_path) -> None:
    """No table is synthesized when no section supplies tabular data (Req 10.1).

    The builder previously fabricated a "Summary" section with a Section/Status
    table to satisfy an unconditional ">=1 table" guarantee. That fabrication is
    removed: when no section provides a ``table`` payload or a body markdown
    table, the document contains ZERO tables and no synthesized "Summary" heading.
    """

    sections = [
        {"heading": "Introduction", "level": 1, "body": "Some plain body text."},
        {"heading": "Details", "level": 2, "body": "More narrative prose."},
    ]

    output_path = tmp_path / "no_table.docx"
    builder = DocumentBuilder("1F4E79")
    written = builder.build(
        title="No Table Report",
        prepared_by="Tester",
        sections=sections,
        output_path=output_path,
    )

    document = Document(str(written))

    # The fabricated Section/Status "Summary" table is gone entirely.
    assert len(document.tables) == 0
    paragraph_texts = [p.text for p in document.paragraphs]
    assert "Summary" not in paragraph_texts


def test_toc_entries_show_page_numbers_with_pageref_fields(tmp_path) -> None:
    """TOC entries render page numbers backed by PAGEREF fields (Req 10.1).

    Each cached TOC entry shows the heading, a dotted leader tab, and a page
    number produced by a ``PAGEREF`` field targeting a bookmark on the matching
    section heading, so Word fills in exact pages on open while a number stays
    visible in non-refreshing viewers.
    """

    sections = [
        {"heading": "First Section", "level": 1, "body": "Body.", "bullets": ["x"]},
        {"heading": "Second Section", "level": 2, "body": "More."},
    ]

    output_path = tmp_path / "toc_pages.docx"
    builder = DocumentBuilder("1F4E79")
    written = builder.build(
        title="Paged Report",
        prepared_by="Tester",
        sections=sections,
        output_path=output_path,
    )

    document = Document(str(written))
    document_xml = document.element.xml

    # PAGEREF fields target the heading bookmarks, and bookmarks are emitted for
    # the section headings so the references resolve.
    assert "PAGEREF" in document_xml
    assert "_Toc_0" in document_xml
    assert "w:bookmarkStart" in document_xml

    # A TOC entry paragraph carries the heading text plus a visible page number.
    toc_entry_texts = [
        p.text
        for p in document.paragraphs
        if "First Section" in p.text and any(ch.isdigit() for ch in p.text)
    ]
    assert toc_entry_texts, "expected a TOC entry with a visible page number"


def test_toc_cached_page_numbers_are_content_length_based(tmp_path) -> None:
    """Cached TOC page numbers reflect content length, not a per-entry increment.

    Several short sections flow continuously with no page breaks between them, so
    their cached TOC page numbers must be monotonic non-decreasing, start at 2
    (content begins on page 2, after the cover + TOC on page 1), and stay small:
    multiple short sections share a page rather than each incrementing the number
    (which the old naive 2, 3, 4, ... scheme produced). PAGEREF fields remain so
    Microsoft Word still fills in the exact page numbers on open.
    """

    # Several short sections, each with a bullet so no synthesized "Key
    # Considerations" entry is appended (keeps entry count == section count).
    sections = [
        {
            "heading": f"Section {n}",
            "level": 1,
            "body": "A short paragraph of body text.",
            "bullets": ["A concise point"],
        }
        for n in range(1, 7)
    ]

    output_path = tmp_path / "estimated_pages.docx"
    builder = DocumentBuilder("1F4E79")
    written = builder.build(
        title="Estimated Pages Report",
        prepared_by="Tester",
        sections=sections,
        output_path=output_path,
    )

    document = Document(str(written))
    document_xml = document.element.xml

    # PAGEREF fields are still present so Word computes exact pages on open.
    assert "PAGEREF" in document_xml

    # Collect the cached page numbers from the TOC entry paragraphs. A TOC entry
    # is a non-heading paragraph carrying a tab followed by the cached number.
    cached_numbers: list[int] = []
    for paragraph in document.paragraphs:
        if paragraph.style is not None and paragraph.style.name.startswith("Heading"):
            continue
        text = paragraph.text
        if "\t" not in text:
            continue
        trailing = text.split("\t")[-1].strip()
        if trailing.isdigit():
            cached_numbers.append(int(trailing))

    # One cached number per section entry.
    assert len(cached_numbers) == len(sections)

    # They start at page 2 and are monotonic non-decreasing.
    assert cached_numbers[0] == 2
    assert all(
        earlier <= later
        for earlier, later in zip(cached_numbers, cached_numbers[1:], strict=False)
    )

    # The numbers stay small: multiple short sections share a page, so the max is
    # far below what a naive one-page-per-section scheme (2, 3, 4, ...) yields.
    naive_max = 2 + (len(sections) - 1)
    assert max(cached_numbers) < naive_max
    assert max(cached_numbers) <= len(sections)


# --- Property 15 ------------------------------------------------------------


# Feature: autonomous-agent-service, Property 15: Theme color resolution is safe and total
@pytest.mark.property
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture],
)
@given(color=st.one_of(_valid_hex(), st.text(max_size=12)))
def test_theme_color_resolution_is_safe_and_total(
    tmp_path_factory,
    color: str,
) -> None:
    """Property 15: theme-color resolution is safe and total.

    **Validates: Requirements 14.3, 14.4, 14.5**

    For any ``THEME_COLOR`` string, the resolved heading color equals the value
    when it is a valid 6-digit hex, otherwise the documented default; building
    always completes without raising, and an invalid value produces a structured
    warning.
    """

    expected = _expected_resolved(color)
    # Determine independently whether the input normalized to a valid hex color.
    normalized = color.strip()
    if normalized.startswith("#"):
        normalized = normalized[1:]
    input_was_valid = len(normalized) == 6 and all(
        c in "0123456789abcdefABCDEF" for c in normalized
    )

    # Capture structured log entries via an injected sink.
    entries: list[str] = []
    logger = StructuredLogger(sink=entries.append)

    builder = DocumentBuilder(color, logger)

    # Resolution equals the value when valid, else the documented default.
    assert builder.theme_color == expected
    # The resolved color is always a usable 6-digit hex (never raises here).
    RGBColor.from_string(builder.theme_color)

    # Building always completes without raising, regardless of the color input.
    output_path = tmp_path_factory.mktemp("theme") / "themed.docx"
    written = builder.build(
        title="Theme Test",
        prepared_by="Tester",
        sections=[{"heading": "Section", "body": "Body"}],
        output_path=output_path,
    )
    assert written.exists()
    Document(str(written))  # re-opens without raising

    # An invalid value produces a structured warning; a valid value does not warn.
    warnings = [
        json.loads(entry)
        for entry in entries
        if json.loads(entry).get("level") == "WARNING"
    ]
    if input_was_valid:
        assert warnings == []
    else:
        assert len(warnings) >= 1
        assert warnings[0]["component"] == "document_builder"



# --- Rich body rendering (markdown -> Word formatting) ----------------------


def _iter_block_items(document):
    """Yield document body children as ('paragraph', p) / ('table', t) tuples.

    Iterating the body in document order lets tests assert the relative order of
    paragraphs, tables, and page breaks without relying on private internals.
    """

    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("paragraph", Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            yield ("table", Table(child, document))


def test_rich_body_renders_markdown_as_word_formatting(tmp_path) -> None:
    """A markdown body renders as real Word formatting, not literal markup.

    The section body carries a bold label, a ``##`` sub-heading, ``*`` bullets,
    and a markdown table. The produced document must contain: no literal ``**``
    in any paragraph text, a Heading 2/3 sub-heading for the ``## Sub Heading``
    line, ``List Bullet`` paragraphs for the items, a run whose text is the bold
    label (marked bold), and an extra real table for the markdown table.
    """

    body = "\n".join(
        [
            "**Bold Label:** intro text with *emphasis*.",
            "",
            "## Sub Heading",
            "",
            "* item one",
            "* item two",
            "",
            "| Feature | Value |",
            "| --- | --- |",
            "| Speed | Fast |",
            "| Cost | Low |",
        ]
    )
    sections = [{"heading": "Overview", "level": 1, "body": body}]

    output_path = tmp_path / "rich.docx"
    builder = DocumentBuilder("1F4E79")
    written = builder.build(
        title="Rich Body",
        prepared_by="Tester",
        sections=sections,
        output_path=output_path,
    )

    document = Document(str(written))

    # No literal markdown asterisks leak into any paragraph text.
    for paragraph in document.paragraphs:
        assert "**" not in paragraph.text

    # The "## Sub Heading" line becomes a Heading 2/3 sub-heading paragraph.
    subheadings = [
        p
        for p in document.paragraphs
        if p.style is not None
        and p.style.name in ("Heading 2", "Heading 3")
        and p.text == "Sub Heading"
    ]
    assert len(subheadings) >= 1

    # The bullet items render as List Bullet paragraphs.
    bullet_texts = [
        p.text
        for p in document.paragraphs
        if p.style is not None and p.style.name == "List Bullet"
    ]
    assert "item one" in bullet_texts
    assert "item two" in bullet_texts

    # The bold label renders as a run marked bold with the markers stripped.
    bold_runs = [
        run.text
        for p in document.paragraphs
        for run in p.runs
        if run.font.bold
    ]
    assert "Bold Label:" in bold_runs

    # The markdown table renders as an additional real Word table containing the
    # parsed header and data cells.
    assert len(document.tables) >= 1
    markdown_tables = [
        table
        for table in document.tables
        if table.rows and [c.text for c in table.rows[0].cells] == ["Feature", "Value"]
    ]
    assert len(markdown_tables) == 1
    rendered = markdown_tables[0]
    data_cells = {c.text for row in rendered.rows[1:] for c in row.cells}
    assert {"Speed", "Fast", "Cost", "Low"} <= data_cells


def test_cover_and_toc_share_first_page_single_page_break(tmp_path) -> None:
    """The cover block and TOC share page 1; content starts after one page break.

    The cover no longer emits its own page break, so exactly one page break (the
    one after the table of contents) precedes the first content section. The
    cover-block paragraphs (title, date, "Prepared by") and the TOC entries must
    all precede that first page break.
    """

    sections = [
        {
            "heading": "First Section",
            "level": 1,
            "body": "Body text.",
            "bullets": ["A point"],
            "table": {"headers": ["Metric"], "rows": [["Value"]]},
        }
    ]

    output_path = tmp_path / "pagination.docx"
    builder = DocumentBuilder("1F4E79")
    written = builder.build(
        title="Pagination Report",
        prepared_by="Tester",
        sections=sections,
        output_path=output_path,
    )

    document = Document(str(written))

    # Exactly one rendered page break exists in the body (the one after the TOC).
    body_xml = document.element.body.xml
    assert body_xml.count('w:type="page"') == 1

    # Collect the ordered text of paragraphs up to (and excluding) the paragraph
    # that carries the page break.
    from docx.oxml.ns import qn

    texts_before_break = []
    saw_break = False
    for kind, item in _iter_block_items(document):
        if kind != "paragraph":
            continue
        if item._p.findall(".//" + qn("w:br")):
            # A paragraph containing a page break marks the boundary; the break
            # itself lives at the end of the last cover/TOC content.
            saw_break = True
            texts_before_break.append(item.text)
            break
        texts_before_break.append(item.text)

    assert saw_break

    # The cover block and the TOC entries all appear before the first content
    # section begins on the next page. TOC entry text now includes a tab and a
    # page number, so match the heading text as a substring.
    assert "Pagination Report" in texts_before_break
    assert any("Prepared by:" in t for t in texts_before_break)
    assert any("Table of Contents" in t for t in texts_before_break)
    assert any("First Section" in t for t in texts_before_break)
