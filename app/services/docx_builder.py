"""Reusable Word (``.docx``) document builder (Req 10, 14.3-14.5).

This module defines :class:`DocumentBuilder`, the reusable component that turns a
set of section payloads into a polished Microsoft Word deliverable via
``python-docx``. It satisfies the ``DocumentBuilderProtocol`` seam declared in
:mod:`app.agent.tools`, exposing a single :meth:`DocumentBuilder.build` method.

Every generated document contains (Req 10.1):

- a **cover block** with the document title (styled in the theme color), the
  generation date (formatted ``Month DD, YYYY``), and a "Prepared by" line,
  followed by a thin divider and the table of contents on the same first page
  (a single page break after the TOC starts the first content section);
- a **table of contents** inserted as a Word ``TOC`` field, with the document
  settings ``w:updateFields`` flag set so Word refreshes the field on open. Each
  cached TOC entry renders as a real TOC line: the heading text, a dotted tab
  leader, and a right-aligned page number backed by a ``PAGEREF`` field that
  targets a bookmark on the corresponding section heading;
- **styled headings** (``Heading 1``/``Heading 2``) whose run font color is the
  resolved theme color;
- **body text** paragraphs;
- a **formatted table** (header row plus one or more data rows, using the
  visible ``Table Grid`` style) whenever a section actually supplies tabular
  data (``table`` payload or a markdown pipe-table in its body). No table is
  synthesized when none is supplied, so a document may legitimately contain zero
  tables;
- at least one **bullet list** (``List Bullet`` style); and
- a **footer** containing a page-number (``PAGE``) field.

Theme-color resolution is **independent of all other configuration** (Req 14.5)
and **total / safe** (Req 14.3, 14.4): a valid 6-digit hex value (with or without
a leading ``#``) is used as-is; any unset or invalid value falls back to the
documented default :data:`DEFAULT_THEME_COLOR`, emits a structured warning through
the injected logger, and never raises.
"""

from __future__ import annotations

import re
from collections.abc import Mapping, Sequence
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.logging import StructuredLogger

# The documented default theme color (a deep professional blue), applied when the
# configured Theme_Color is unset or invalid (Req 14.4).
DEFAULT_THEME_COLOR = "1F4E79"

# The set of valid hexadecimal digits used to validate a theme-color string.
_HEX_DIGITS = frozenset("0123456789abcdefABCDEF")

# The number of hex digits in a valid 6-digit color.
_HEX_LENGTH = 6

# The date format used on the cover page ("Month DD, YYYY", e.g. "January 05, 2025").
_DATE_FORMAT = "%B %d, %Y"

# The Word style names relied upon for structural guarantees.
_TABLE_STYLE = "Table Grid"
_BULLET_STYLE = "List Bullet"
_NUMBER_STYLE = "List Number"

# The bullet markers recognized at the start of a body line.
_BULLET_MARKERS = ("*", "-", "\u2022")

# Matches a numbered-list line, e.g. "1. First item" -> content "First item".
_NUMBERED_RE = re.compile(r"^\d+\.\s+(?P<content>.*)$")

# Matches a markdown-table separator cell, e.g. "---", ":--", ":-:" (dashes with
# optional leading/trailing colons).
_TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-{1,}:?$")

# Matches a ``**bold**`` span (non-greedy) for inline markdown parsing.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Matches an ``*italic*`` or ``_italic_`` span (non-greedy) for inline parsing.
_ITALIC_RE = re.compile(r"\*(.+?)\*|_(.+?)_")

# The TOC field instruction (headings levels 1-3) inserted for the table of contents.
_TOC_INSTRUCTION = ' TOC \\o "1-3" \\h \\z \\u '

# Approximate number of rendered body characters that fit on a single page. Used
# only to compute a realistic *cached* TOC page number for each entry; Microsoft
# Word recomputes the exact page numbers on open (``w:updateFields``). Chosen in
# the 2600-3000 range for a typical single-column business page.
_CHARS_PER_PAGE = 2800

# Rough character overhead attributed to a rendered heading line (the heading
# text plus its surrounding paragraph spacing) when estimating page positions.
_HEADING_CHAR_OVERHEAD = 60

# Rough characters contributed by a single rendered table row when estimating.
_TABLE_ROW_CHARS = 90

# A small fixed per-section overhead (paragraph spacing, etc.) in the estimate.
_SECTION_CHAR_OVERHEAD = 40

# Estimated rendered characters of the synthesized "Key Considerations" fallback
# bullet section (its heading is counted separately); a small fixed value since
# the fallback emits a short, fixed set of generic bullet points.
_SYNTHESIZED_BULLETS_CHARS = 150


def _is_valid_hex_color(candidate: str) -> bool:
    """Return whether ``candidate`` is exactly six hexadecimal digits.

    Args:
        candidate: A color string with any surrounding whitespace and any single
            leading ``#`` already removed.

    Returns:
        ``True`` when ``candidate`` is a 6-digit hexadecimal string.
    """

    return len(candidate) == _HEX_LENGTH and all(c in _HEX_DIGITS for c in candidate)


def _normalize_hex_color(value: str) -> str | None:
    """Normalize ``value`` to a canonical 6-digit uppercase hex color.

    Strips surrounding whitespace and a single optional leading ``#``. When the
    remaining text is a valid 6-digit hex color, its uppercase form is returned;
    otherwise ``None`` signals an invalid value.

    Args:
        value: The raw theme-color string to normalize.

    Returns:
        The canonical uppercase 6-digit hex string, or ``None`` when invalid.
    """

    candidate = value.strip()
    if candidate.startswith("#"):
        candidate = candidate[1:]
    if _is_valid_hex_color(candidate):
        return candidate.upper()
    return None


class DocumentBuilder:
    """Assemble polished ``.docx`` deliverables with ``python-docx`` (Req 10).

    The builder is reusable: a single instance resolves its theme color once at
    construction and can produce any number of documents through :meth:`build`.

    Attributes:
        theme_color: The resolved, canonical 6-digit uppercase hex color applied
            to styled headings and the cover-page title.
    """

    def __init__(
        self,
        theme_color: str,
        logger: StructuredLogger | None = None,
    ) -> None:
        """Resolve the theme color independently and safely (Req 14.3-14.5).

        The provided ``theme_color`` is resolved on its own, independent of all
        other configuration (Req 14.5). A valid 6-digit hex value (with or
        without a leading ``#``) is adopted as the resolved color; any unset or
        invalid value falls back to :data:`DEFAULT_THEME_COLOR`, emits a
        structured warning through ``logger`` (when provided), and never raises
        (Req 14.4).

        Args:
            theme_color: The configured Theme_Color value to resolve.
            logger: An optional structured logger used to emit a warning when the
                configured color is invalid. When ``None``, no warning is emitted
                but resolution still falls back safely.
        """

        self._logger = logger
        # Monotonic counter giving each heading bookmark a unique ``w:id``; reset
        # at the start of every :meth:`build` since each build is a fresh document.
        self._bookmark_id = 0
        normalized = _normalize_hex_color(theme_color) if isinstance(theme_color, str) else None
        if normalized is None:
            if self._logger is not None:
                self._logger.decision(
                    component="document_builder",
                    run_id="-",
                    decision="invalid or unset THEME_COLOR; applying documented default",
                    level="WARNING",
                    invalid_value=theme_color,
                    default_theme_color=DEFAULT_THEME_COLOR,
                )
            self.theme_color: str = DEFAULT_THEME_COLOR
        else:
            self.theme_color = normalized

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def build(
        self,
        *,
        title: str,
        prepared_by: str,
        sections: Sequence[Mapping[str, Any]],
        output_path: Path,
    ) -> Path:
        """Assemble a ``.docx`` deliverable and return the written path (Req 10.1).

        The produced document always contains a cover page, a table-of-contents
        field, at least one styled heading, at least one bullet list, and a
        footer page-number field. A formatted table is rendered only when a
        section actually supplies tabular data (a ``table`` payload or a markdown
        pipe-table in its body); no table is fabricated when none is supplied, so
        the document may legitimately contain zero tables. Missing keys in a
        section mapping are tolerated, and the bullet-list element is synthesized
        when the supplied sections do not provide one, so that structural
        guarantee holds for any input.

        Each cached TOC entry is backed by a ``PAGEREF`` field targeting a
        bookmark placed on the corresponding rendered section heading; the TOC
        entries and bookmarked headings are emitted in the same order and count
        so the page references line up.

        Args:
            title: The document title rendered on the cover page.
            prepared_by: The cover-page "Prepared by" line value.
            sections: The ordered section payloads to render. Each mapping may
                provide ``heading`` (str), ``level`` (1 or 2), ``body`` (str),
                ``bullets`` (list of str), and ``table``
                (``{"headers": [...], "rows": [[...]]}``); all keys are optional.
            output_path: The filesystem path the document is written to; parent
                directories are created as needed.

        Returns:
            The path of the written ``.docx`` artifact (equal to ``output_path``).
        """

        document = Document()

        # Each build produces a fresh document, so reset the bookmark id counter
        # used to give every section-heading bookmark a unique ``w:id``.
        self._bookmark_id = 0

        # Compute the full ordered list of headings the document will render
        # (real sections plus any synthesized "Key Considerations" section)
        # BEFORE the TOC is written, so the table of contents lists the actual
        # document structure rather than a placeholder (Req 10.1). Each entry's
        # position is its bookmark index, so section headings can be tagged with
        # matching bookmarks that the TOC ``PAGEREF`` fields target.
        toc_entries = self._compute_toc_entries(sections)
        # Estimate a realistic starting page number for each TOC entry from the
        # accumulated rendered content before it (content flows continuously with
        # no per-section page breaks), used as the cached PAGEREF display text.
        toc_pages = self._estimate_page_numbers(sections)

        self._add_cover_page(document, title=title, prepared_by=prepared_by)
        self._add_table_of_contents(document, toc_entries, toc_pages)

        has_bullets = False
        for index, section in enumerate(sections):
            _rendered_table, rendered_bullets = self._add_section(
                document, section, bookmark_name=f"_Toc_{index}"
            )
            has_bullets = has_bullets or rendered_bullets

        # Guarantee at least one bullet list exists overall (Req 10.1). A table
        # is intentionally NOT synthesized: it is rendered only when a section
        # supplies real tabular data (owner-directed relaxation of the earlier
        # unconditional ">=1 table" guarantee).
        if not has_bullets:
            self._add_default_bullets(
                document, sections, bookmark_name=f"_Toc_{len(sections)}"
            )

        self._add_footer_page_number(document)
        self._enable_update_fields(document)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Cover page / TOC
    # ------------------------------------------------------------------

    def _add_cover_page(
        self, document: Document, *, title: str, prepared_by: str
    ) -> None:
        """Render the cover block (title, date, "Prepared by") on page 1 (Req 10.1).

        The cover no longer emits its own trailing page break. Instead it adds a
        small centered divider/spacer so the cover block and the table of
        contents (rendered next) sit together on the first page. The single page
        break that starts the first content section on a fresh page is emitted
        after the table of contents (see :meth:`_add_table_of_contents`).

        Args:
            document: The document being assembled.
            title: The document title, styled in the theme color.
            prepared_by: The value for the "Prepared by" line.
        """

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor.from_string(self.theme_color)

        date_paragraph = document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(datetime.now().strftime(_DATE_FORMAT))

        prepared_paragraph = document.add_paragraph()
        prepared_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prepared_paragraph.add_run(f"Prepared by: {prepared_by}")

        # A thin, tasteful divider separating the cover block from the TOC so the
        # two share page 1 (rather than the cover consuming an entire page).
        divider_paragraph = document.add_paragraph()
        divider_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider_run = divider_paragraph.add_run("\u2014\u2014\u2014")
        divider_run.font.color.rgb = RGBColor.from_string(self.theme_color)

    def _compute_toc_entries(
        self, sections: Sequence[Mapping[str, Any]]
    ) -> list[tuple[str, int]]:
        """Compute the ordered ``(heading, level)`` list the document will render.

        The result mirrors exactly what :meth:`build` renders: one entry per
        supplied section (its heading text and normalized level), followed by the
        synthesized ``Key Considerations`` (level 2) section when no section
        provides a bullet list. No ``Summary`` entry is synthesized: a table is
        rendered only when a section supplies real tabular data, never
        fabricated. Computing this list up front lets the table of contents list
        the real document structure (Req 10.1) and lets each entry's position
        serve as the bookmark index shared with the rendered section headings.

        Args:
            sections: The ordered section payloads that will be rendered.

        Returns:
            The ordered list of ``(heading_text, level)`` tuples, where ``level``
            is always ``1`` or ``2``.
        """

        entries: list[tuple[str, int]] = []
        has_bullets = False
        for section in sections:
            heading = str(section.get("heading", "") or "")
            level = section.get("level", 1)
            if level not in (1, 2):
                level = 1
            entries.append((heading, level))

            bullets = section.get("bullets")
            if isinstance(bullets, (list, tuple)) and bullets:
                has_bullets = True

        # This mirrors the synthesized bullet section build() appends to
        # guarantee a bullet list always exists (Req 10.1). No table entry is
        # synthesized -- tables are rendered only from real data.
        if not has_bullets:
            entries.append(("Key Considerations", 2))
        return entries

    def _estimate_page_numbers(
        self, sections: Sequence[Mapping[str, Any]]
    ) -> list[int]:
        """Estimate a realistic starting page number for each TOC entry (Req 10.1).

        The estimate walks the SAME ordered entries as :meth:`_compute_toc_entries`
        (the real sections plus the synthesized "Key Considerations" fallback when
        present) and assigns each entry a starting page from the accumulated
        rendered content BEFORE it. Because content sections flow continuously
        (there is no page break between them — only one break after the TOC),
        multiple short sections legitimately share a page, so this replaces the
        old naive per-entry increment (2, 3, 4, ...) with a content-length model:

        - Content begins on page 2 (page 1 holds the cover block and the TOC).
        - Each section's rendered size (in characters) is estimated as the heading
          overhead plus its body length, its bullet lengths, its table rows
          (``rows * _TABLE_ROW_CHARS``), and a small fixed per-section overhead.
        - A page holds roughly :data:`_CHARS_PER_PAGE` characters, so an entry's
          start page is ``2 + cumulative_chars_before_it // _CHARS_PER_PAGE``. The
          numbers are therefore monotonic non-decreasing.

        The estimate is defensive: it never raises, tolerates missing or oversized
        fields, and caps each number so the last entry cannot exceed ``2 +
        total_estimated_pages``. On any unexpected error it falls back to the
        monotonic ``index + 2`` estimate.

        Args:
            sections: The ordered section payloads that will be rendered.

        Returns:
            One integer per TOC entry, aligned by index with
            :meth:`_compute_toc_entries`.
        """

        entries = self._compute_toc_entries(sections)
        try:
            sizes: list[int] = []
            for index, entry in enumerate(entries):
                section = sections[index] if index < len(sections) else None
                sizes.append(self._estimate_section_chars(entry, section))

            total_chars = sum(sizes)
            # ``ceil`` of the total content over the per-page capacity, at least 1.
            total_pages = max(1, -(-total_chars // _CHARS_PER_PAGE))
            max_page = 2 + total_pages

            numbers: list[int] = []
            cumulative = 0
            for size in sizes:
                page = 2 + cumulative // _CHARS_PER_PAGE
                if page > max_page:
                    page = max_page
                numbers.append(page)
                cumulative += size
            return numbers
        except Exception:  # noqa: BLE001 - estimation must never fail a build
            return [index + 2 for index in range(len(entries))]

    @staticmethod
    def _estimate_section_chars(
        entry: tuple[str, int], section: Mapping[str, Any] | None
    ) -> int:
        """Estimate the rendered character count of one TOC entry's section.

        Args:
            entry: The ``(heading, level)`` tuple for the entry (its heading text
                contributes to the estimate).
            section: The section payload backing the entry, or ``None`` for the
                synthesized "Key Considerations" fallback bullet section.

        Returns:
            A non-negative estimated character count; a small default on any
            unexpected error so estimation never raises.
        """

        try:
            heading_text, _level = entry
            chars = _HEADING_CHAR_OVERHEAD + len(str(heading_text or ""))

            if section is None:
                # The synthesized "Key Considerations" fallback bullet section.
                chars += _SYNTHESIZED_BULLETS_CHARS
            else:
                body = section.get("body")
                if isinstance(body, str):
                    chars += len(body)

                bullets = section.get("bullets")
                if isinstance(bullets, (list, tuple)):
                    for bullet in bullets:
                        chars += len(str(bullet))

                table = section.get("table")
                if isinstance(table, Mapping):
                    rows = table.get("rows")
                    if isinstance(rows, (list, tuple)):
                        chars += len(rows) * _TABLE_ROW_CHARS

            chars += _SECTION_CHAR_OVERHEAD
            return max(0, chars)
        except Exception:  # noqa: BLE001 - estimation must never fail a build
            return _HEADING_CHAR_OVERHEAD + _SECTION_CHAR_OVERHEAD

    def _add_table_of_contents(
        self,
        document: Document,
        entries: Sequence[tuple[str, int]],
        page_numbers: Sequence[int] | None = None,
    ) -> None:
        """Insert a heading and a real, visible Word ``TOC`` field (Req 10.1).

        The TOC is written as a proper complex field (``fldChar`` begin /
        ``instrText`` / ``fldChar`` separate / cached result / ``fldChar`` end).
        The cached "result" content is one visible paragraph per heading rendered
        as a real TOC line: the heading text on the left, a dotted tab leader, and
        a right-aligned page number produced by a ``PAGEREF`` complex field that
        targets a bookmark on the corresponding section heading. Level-2 entries
        are indented. Each ``PAGEREF`` also carries a cached, best-effort page
        number so a number is visible even in viewers that do not refresh fields;
        the ``w:updateFields`` settings flag (set in :meth:`build`) instructs Word
        to replace the cached numbers with exact pages and rebuild the real,
        page-numbered TOC when the document is opened.

        The entries are emitted in the same order as the bookmarked section
        headings, so entry ``i`` targets bookmark ``_Toc_{i}`` (Req 10.1).

        Args:
            document: The document being assembled.
            entries: The ordered ``(heading, level)`` list to render as the
                cached TOC result.
            page_numbers: The per-entry cached starting page numbers (aligned by
                index with ``entries``). When ``None`` or shorter than
                ``entries``, entries fall back to a monotonic ``index + 2``
                estimate.
        """

        self._add_styled_heading(document, "Table of Contents", level=1)

        # Open the complex field: begin -> instruction -> separate.
        field_paragraph = document.add_paragraph()
        self._append_fld_char(field_paragraph, "begin")
        self._append_instr_text(field_paragraph, _TOC_INSTRUCTION)
        self._append_fld_char(field_paragraph, "separate")

        # Cached result: one visible paragraph per heading so the TOC is readable
        # in any viewer. Each line shows the heading, a dotted leader tab, and a
        # PAGEREF-backed page number. The closing ``end`` fldChar is appended to
        # the last rendered paragraph (or the opening paragraph when there are no
        # entries).
        last_paragraph = field_paragraph
        for index, (text, level) in enumerate(entries):
            entry_paragraph = document.add_paragraph()
            if level == 2:
                entry_paragraph.paragraph_format.left_indent = Pt(18)
            # A right-aligned tab stop with a dotted leader near the right margin
            # produces the classic "heading .... page" TOC line.
            entry_paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            entry_paragraph.add_run(text)
            entry_paragraph.add_run("\t")
            # Content starts on page 2 (cover + TOC share page 1). The cached page
            # number is a content-length-based estimate (multiple short sections
            # may share a page) rather than a naive per-entry increment; Word
            # overwrites these with the exact page numbers on open.
            if page_numbers is not None and index < len(page_numbers):
                cached_text = str(page_numbers[index])
            else:
                cached_text = str(index + 2)
            self._append_pageref_field(
                entry_paragraph,
                bookmark_name=f"_Toc_{index}",
                cached_text=cached_text,
            )
            last_paragraph = entry_paragraph

        self._append_fld_char(last_paragraph, "end")

        document.add_page_break()

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------

    def _add_section(
        self,
        document: Document,
        section: Mapping[str, Any],
        *,
        bookmark_name: str | None = None,
    ) -> tuple[bool, bool]:
        """Render a single section, tolerating missing keys.

        Args:
            document: The document being assembled.
            section: The section payload mapping.
            bookmark_name: The bookmark name to wrap the section heading with so a
                TOC ``PAGEREF`` field can target it. When ``None``, no bookmark is
                emitted.

        Returns:
            A ``(rendered_table, rendered_bullets)`` tuple indicating whether this
            section contributed a formatted table and/or a bullet list.
        """

        heading = str(section.get("heading", "") or "")
        level = section.get("level", 1)
        if level not in (1, 2):
            level = 1
        self._add_styled_heading(
            document, heading, level=level, bookmark_name=bookmark_name
        )

        body = section.get("body")
        if body:
            self._add_rich_body(document, str(body))

        rendered_bullets = False
        bullets = section.get("bullets")
        if isinstance(bullets, (list, tuple)) and bullets:
            for bullet in bullets:
                document.add_paragraph(str(bullet), style=_BULLET_STYLE)
            rendered_bullets = True

        rendered_table = False
        table = section.get("table")
        if isinstance(table, Mapping):
            headers = table.get("headers")
            rows = table.get("rows")
            if (
                isinstance(headers, (list, tuple))
                and headers
                and isinstance(rows, (list, tuple))
                and rows
            ):
                self._add_table(document, headers, rows)
                rendered_table = True

        return rendered_table, rendered_bullets

    # ------------------------------------------------------------------
    # Rich body rendering (markdown -> real Word formatting)
    # ------------------------------------------------------------------

    def _add_rich_body(self, document: Document, text: str) -> None:
        """Render a markdown-ish body string as real Word formatting.

        The body returned by the LLM commonly contains lightweight markdown:
        ``#`` headings, ``*``/``-``/``•`` bullets, ``1.`` numbered items, pipe
        tables, and inline ``**bold**``/``*italic*`` spans. Inserting that verbatim
        as one plain paragraph exposes raw markup to the reader. This helper parses
        the text line-by-line and renders each block with the appropriate Word
        construct instead:

        - a line of 1-4 leading ``#`` characters becomes a styled sub-heading
          (``##`` and below map to Heading 2, ``###``/``####`` to Heading 3);
        - ``*``/``-``/``•`` bullet lines become ``List Bullet`` paragraphs;
        - ``<n>.`` lines become ``List Number`` paragraphs (falling back to
          ``List Bullet`` when the numbered style is unavailable);
        - a contiguous block of pipe-delimited rows with a dashes separator row
          becomes a real Word table via :meth:`_add_table`; and
        - any other non-empty line becomes a body paragraph with inline
          ``**bold**`` / ``*italic*`` markers converted to run formatting.

        The parser is fully self-contained (no new dependencies), tolerates
        arbitrary text, and never raises: anything that does not match a
        recognized block is rendered as a plain paragraph.

        Args:
            document: The document being assembled.
            text: The raw body text to render.
        """

        lines = text.split("\n")
        total = len(lines)
        index = 0
        while index < total:
            stripped = lines[index].strip()

            # Blank lines separate blocks and carry no content.
            if not stripped:
                index += 1
                continue

            # Markdown table: a contiguous run of pipe-delimited rows.
            if stripped.startswith("|") and stripped.endswith("|"):
                block: list[str] = []
                cursor = index
                while cursor < total:
                    candidate = lines[cursor].strip()
                    if candidate.startswith("|") and candidate.endswith("|"):
                        block.append(candidate)
                        cursor += 1
                    else:
                        break
                if len(block) >= 2 and self._try_add_markdown_table(document, block):
                    index = cursor
                    continue
                # Not a valid table; fall through and treat this line as a paragraph.

            # Markdown sub-heading (## / ### ...).
            heading_level = self._markdown_heading_level(stripped)
            if heading_level is not None:
                content = self._strip_wrapping_bold(stripped.lstrip("#").strip())
                self._add_styled_heading(document, content, level=heading_level)
                index += 1
                continue

            # Bullet line.
            if self._is_bullet_line(stripped):
                content = stripped[1:].strip()
                paragraph = self._add_body_paragraph(document, style=_BULLET_STYLE)
                self._add_inline_runs(paragraph, content)
                index += 1
                continue

            # Numbered line.
            numbered = _NUMBERED_RE.match(stripped)
            if numbered is not None:
                paragraph = self._add_body_paragraph(document, style=_NUMBER_STYLE)
                self._add_inline_runs(paragraph, numbered.group("content").strip())
                index += 1
                continue

            # Normal paragraph with inline markdown.
            paragraph = document.add_paragraph()
            self._add_inline_runs(paragraph, self._strip_wrapping_bold(stripped))
            index += 1

    @staticmethod
    def _markdown_heading_level(line: str) -> int | None:
        """Return the sub-heading level for a markdown heading line, else ``None``.

        A line beginning with one to four ``#`` characters is treated as a
        sub-heading. ``#``/``##`` map to Heading level 2 and ``###``/``####`` to
        Heading level 3 so that section-level headings remain dominant (Heading 1
        is never produced from body markdown).

        Args:
            line: The already-stripped body line.

        Returns:
            ``2`` or ``3`` for a recognized heading, or ``None`` otherwise.
        """

        hash_count = len(line) - len(line.lstrip("#"))
        if not 1 <= hash_count <= 4:
            return None
        remainder = line[hash_count:]
        # Require the hashes to be followed by whitespace (or be the whole line),
        # so tokens like "#tag" are not misread as headings.
        if remainder and not remainder[0].isspace():
            return None
        return 2 if hash_count <= 2 else 3

    @staticmethod
    def _is_bullet_line(line: str) -> bool:
        """Return whether ``line`` starts with a bullet marker followed by a space.

        Args:
            line: The already-stripped body line.

        Returns:
            ``True`` when the line begins with ``*``, ``-`` or ``•`` and a space
            (so ``**bold**`` intro lines are not mistaken for bullets).
        """

        return len(line) >= 2 and line[0] in _BULLET_MARKERS and line[1] == " "

    def _add_body_paragraph(self, document: Document, *, style: str) -> Any:
        """Add an empty body paragraph in ``style``, falling back to bullets.

        Args:
            document: The document being assembled.
            style: The desired Word paragraph style name.

        Returns:
            The created paragraph. If ``style`` is unavailable in the template,
            the paragraph is created with the ``List Bullet`` style instead.
        """

        try:
            return document.add_paragraph(style=style)
        except KeyError:
            return document.add_paragraph(style=_BULLET_STYLE)

    def _try_add_markdown_table(self, document: Document, block: Sequence[str]) -> bool:
        """Attempt to render a markdown pipe-table ``block`` as a Word table.

        The block must be a header row, a separator row of dashes/colons, and
        zero or more data rows. Cells are parsed by splitting on ``|`` and
        trimming; the separator row is dropped. When the block does not look like
        a valid table (no separator row, or no header cells) the method makes no
        changes and returns ``False`` so the caller can fall back to paragraphs.

        Args:
            document: The document being assembled.
            block: The contiguous pipe-delimited lines (each already stripped and
                starting/ending with ``|``).

        Returns:
            ``True`` when a table was rendered, ``False`` otherwise.
        """

        rows = [self._parse_table_row(line) for line in block]
        if len(rows) < 2 or not self._is_separator_row(rows[1]):
            return False
        headers = rows[0]
        if not headers:
            return False
        data_rows = rows[2:]
        self._add_table(document, headers, data_rows)
        return True

    @staticmethod
    def _parse_table_row(line: str) -> list[str]:
        """Split a pipe-delimited table row into trimmed cell values.

        Args:
            line: A row that starts and ends with ``|`` (e.g. ``"| a | b |"``).

        Returns:
            The trimmed inner cell values (leading/trailing empty splits removed).
        """

        return [cell.strip() for cell in line.split("|")[1:-1]]

    @staticmethod
    def _is_separator_row(cells: Sequence[str]) -> bool:
        """Return whether ``cells`` form a markdown table separator row.

        Args:
            cells: The parsed cell values of a candidate separator row.

        Returns:
            ``True`` when every cell matches a dashes/colons pattern (e.g.
            ``---``, ``:--``, ``:-:``) and there is at least one cell.
        """

        return bool(cells) and all(
            _TABLE_SEPARATOR_CELL_RE.fullmatch(cell) is not None for cell in cells
        )

    @staticmethod
    def _strip_wrapping_bold(text: str) -> str:
        """Strip a single pair of ``**`` markers that wrap the whole ``text``.

        This handles lines like ``"**Key Benefits:**"`` used as pseudo-headings,
        leaving inner ``**`` spans (handled by the inline parser) untouched.

        Args:
            text: The text to unwrap.

        Returns:
            ``text`` without a single enclosing ``**...**`` pair, when present.
        """

        candidate = text.strip()
        if len(candidate) >= 4 and candidate.startswith("**") and candidate.endswith("**"):
            return candidate[2:-2].strip()
        return text

    def _add_inline_runs(self, paragraph: Any, text: str) -> None:
        """Add ``text`` to ``paragraph`` converting inline ``**bold``/``*italic*``.

        Balanced ``**...**`` spans become bold runs and ``*...*``/``_..._`` spans
        become italic runs; the markers themselves are stripped. Text outside any
        span is added as plain runs. When ``text`` contains no markers it is added
        as a single plain run.

        Args:
            paragraph: The paragraph to append runs to.
            text: The (block-level markers already removed) text to render.
        """

        for segment, bold, italic in self._parse_inline(text):
            if segment == "":
                continue
            run = paragraph.add_run(segment)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True

    def _parse_inline(self, text: str) -> list[tuple[str, bool, bool]]:
        """Parse inline markdown into ``(segment, bold, italic)`` tuples.

        The text is first split on ``**bold**`` spans; each resulting fragment is
        then split on ``*italic*``/``_italic_`` spans. Bold spans may themselves
        contain italic spans.

        Args:
            text: The text to parse.

        Returns:
            The ordered list of ``(segment_text, is_bold, is_italic)`` tuples.
        """

        segments: list[tuple[str, bool, bool]] = []
        position = 0
        for match in _BOLD_RE.finditer(text):
            if match.start() > position:
                for piece, italic in self._split_italic(text[position : match.start()]):
                    segments.append((piece, False, italic))
            for piece, italic in self._split_italic(match.group(1)):
                segments.append((piece, True, italic))
            position = match.end()
        if position < len(text):
            for piece, italic in self._split_italic(text[position:]):
                segments.append((piece, False, italic))
        return segments

    @staticmethod
    def _split_italic(text: str) -> list[tuple[str, bool]]:
        """Split ``text`` on ``*italic*``/``_italic_`` spans.

        Args:
            text: The (bold-free) fragment to split.

        Returns:
            An ordered list of ``(segment_text, is_italic)`` tuples.
        """

        pieces: list[tuple[str, bool]] = []
        position = 0
        for match in _ITALIC_RE.finditer(text):
            if match.start() > position:
                pieces.append((text[position : match.start()], False))
            inner = match.group(1) if match.group(1) is not None else match.group(2)
            pieces.append((inner, True))
            position = match.end()
        if position < len(text):
            pieces.append((text[position:], False))
        return pieces

    # ------------------------------------------------------------------
    # Structural helpers
    # ------------------------------------------------------------------

    def _add_styled_heading(
        self,
        document: Document,
        text: str,
        *,
        level: int,
        bookmark_name: str | None = None,
    ) -> None:
        """Add a ``Heading {level}`` whose runs use the theme color (Req 10.1).

        When ``bookmark_name`` is provided, the heading's content is wrapped in a
        Word bookmark so a table-of-contents ``PAGEREF`` field can target it. Only
        real document section headings (those that appear as TOC entries) are
        bookmarked; the "Table of Contents" heading itself and in-body markdown
        sub-headings pass ``None`` and are not bookmarked.

        Args:
            document: The document being assembled.
            text: The heading text.
            level: The heading level (1 or 2).
            bookmark_name: The unique, stable bookmark name to wrap the heading
                with, or ``None`` to emit no bookmark.
        """

        heading = document.add_heading(text, level=level)
        theme_rgb = RGBColor.from_string(self.theme_color)
        for run in heading.runs:
            run.font.color.rgb = theme_rgb
        if bookmark_name is not None:
            self._add_bookmark(heading, bookmark_name)

    def _add_table(
        self,
        document: Document,
        headers: Sequence[Any],
        rows: Sequence[Sequence[Any]],
    ) -> None:
        """Add a bordered table with a header row and the supplied data rows.

        Args:
            document: The document being assembled.
            headers: The column header cells.
            rows: The data rows; extra cells beyond the header width are ignored.
        """

        column_count = len(headers)
        table = document.add_table(rows=1, cols=column_count)
        table.style = _TABLE_STYLE

        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = str(header)

        for row in rows:
            cells = table.add_row().cells
            for index in range(column_count):
                value = row[index] if index < len(row) else ""
                cells[index].text = str(value)

    def _add_default_bullets(
        self,
        document: Document,
        sections: Sequence[Mapping[str, Any]],
        *,
        bookmark_name: str | None = None,
    ) -> None:
        """Synthesize a bullet list so a bullet list always exists (Req 10.1).

        This fallback runs only when no supplied section provided a bullet list.
        It emits a small, fixed set of generic, professional "Key Considerations"
        points rather than echoing the section headings (which would merely
        duplicate the table of contents), while still guaranteeing at least one
        bullet-list paragraph exists (Property 13).

        Args:
            document: The document being assembled.
            sections: The section payloads (unused; the fallback is a fixed,
                sensible default independent of the section headings).
            bookmark_name: The bookmark name to wrap the "Key Considerations"
                heading with so its TOC ``PAGEREF`` entry can target it.
        """

        del sections  # The default bullets are intentionally content-independent.
        self._add_styled_heading(
            document, "Key Considerations", level=2, bookmark_name=bookmark_name
        )
        points = [
            "Objectives and scope are defined above.",
            "Recommendations are detailed in the sections above.",
            "Next steps and timeline to be confirmed with stakeholders.",
        ]
        for point in points:
            document.add_paragraph(point, style=_BULLET_STYLE)

    def _add_footer_page_number(self, document: Document) -> None:
        """Add a centered ``PAGE`` field to the primary section footer (Req 10.1).

        Args:
            document: The document being assembled.
        """

        footer = document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("Page ")
        self._add_field(paragraph, instruction=" PAGE ", cached_text="1")

    # ------------------------------------------------------------------
    # Low-level Word XML helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _add_field(paragraph: Any, *, instruction: str, cached_text: str = "") -> None:
        """Append a simple Word field (``w:fldSimple``) to ``paragraph``.

        Using a simple field lets ``python-docx`` embed a dynamic field (such as
        the footer ``PAGE`` number) that Word evaluates when the document is
        opened. The multi-paragraph TOC uses a complex field instead (see
        :meth:`_append_fld_char`).

        Args:
            paragraph: The paragraph to append the field to.
            instruction: The field instruction (e.g. ``" PAGE "`` or a TOC spec).
            cached_text: Placeholder text shown until Word refreshes the field.
        """

        fld_simple = OxmlElement("w:fldSimple")
        fld_simple.set(qn("w:instr"), instruction)
        run = OxmlElement("w:r")
        text_element = OxmlElement("w:t")
        text_element.text = cached_text
        run.append(text_element)
        fld_simple.append(run)
        paragraph._p.append(fld_simple)

    @staticmethod
    def _append_fld_char(paragraph: Any, fld_char_type: str) -> None:
        """Append a complex-field ``w:fldChar`` run to ``paragraph``.

        Complex fields (unlike ``w:fldSimple``) are delimited by ``begin``,
        ``separate``, and ``end`` field characters, which lets a field's cached
        result span its own runs and paragraphs — used here so the TOC result can
        be a readable, multi-paragraph list of headings.

        Args:
            paragraph: The paragraph to append the field-character run to.
            fld_char_type: The field-character type (``"begin"``, ``"separate"``
                or ``"end"``).
        """

        run = OxmlElement("w:r")
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), fld_char_type)
        run.append(fld_char)
        paragraph._p.append(run)

    def _add_bookmark(self, paragraph: Any, name: str) -> None:
        """Wrap ``paragraph``'s content in a Word bookmark named ``name``.

        Emits a ``w:bookmarkStart`` (carrying a unique, incrementing ``w:id`` and
        the given ``name``) immediately after the paragraph properties and a
        matching ``w:bookmarkEnd`` at the end of the paragraph, so the bookmark
        brackets the heading's runs. A TOC ``PAGEREF`` field targeting ``name``
        then resolves to the bookmarked heading's page.

        Args:
            paragraph: The heading paragraph to bookmark.
            name: The unique, stable bookmark name (e.g. ``"_Toc_0"``).
        """

        bookmark_id = str(self._bookmark_id)
        self._bookmark_id += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), name)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)

        paragraph_element = paragraph._p
        properties = paragraph_element.find(qn("w:pPr"))
        if properties is not None:
            properties.addnext(start)
        else:
            paragraph_element.insert(0, start)
        paragraph_element.append(end)

    def _append_pageref_field(
        self, paragraph: Any, *, bookmark_name: str, cached_text: str
    ) -> None:
        """Append a ``PAGEREF`` complex field targeting ``bookmark_name``.

        The field is written as ``fldChar`` begin -> ``instrText``
        (`` PAGEREF <name> \\h ``) -> ``fldChar`` separate -> a cached page-number
        text run -> ``fldChar`` end. Word replaces the cached number with the
        exact page of the bookmarked heading when the document is opened (the
        ``w:updateFields`` flag is set), while the cached number keeps a value
        visible in viewers that do not refresh fields.

        Args:
            paragraph: The TOC entry paragraph to append the field to.
            bookmark_name: The name of the bookmark on the target heading.
            cached_text: The best-effort cached page number shown until refresh.
        """

        self._append_fld_char(paragraph, "begin")
        self._append_instr_text(paragraph, f" PAGEREF {bookmark_name} \\h ")
        self._append_fld_char(paragraph, "separate")
        paragraph.add_run(cached_text)
        self._append_fld_char(paragraph, "end")

    @staticmethod
    def _append_instr_text(paragraph: Any, instruction: str) -> None:
        """Append a ``w:instrText`` run carrying a complex-field instruction.

        Args:
            paragraph: The paragraph to append the instruction run to.
            instruction: The field instruction text (e.g. the ``TOC`` spec). It
                is written with ``xml:space="preserve"`` so its leading and
                trailing spaces are retained.
        """

        run = OxmlElement("w:r")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction
        run.append(instr_text)
        paragraph._p.append(run)

    @staticmethod
    def _enable_update_fields(document: Document) -> None:
        """Set the ``w:updateFields`` settings flag so Word refreshes fields.

        This instructs Word to update dynamic fields (notably the TOC) the first
        time the document is opened, since ``python-docx`` does not compute field
        results itself.

        Args:
            document: The document being assembled.
        """

        settings_element = document.settings.element
        update_fields = settings_element.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings_element.append(update_fields)
        update_fields.set(qn("w:val"), "true")
